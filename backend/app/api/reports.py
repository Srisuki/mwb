from __future__ import annotations

import csv
import uuid
from datetime import datetime, timezone
from io import BytesIO, StringIO

from docx import Document as WordDocument
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func, or_, select
from sqlalchemy.orm import Session

from app.api.deps import ensure_entity_access, permitted_entity_ids, require_permission
from app.db.session import get_db
from app.models import (
    AppSetting,
    AuditArea,
    ChecklistItem,
    Entity,
    Observation,
    Report,
    ReportLock,
    ReportSequence,
    ReportStatus,
    User,
)
from app.schemas import ReportGenerate, ReportOut
from app.services.audit import record_audit

router = APIRouter(prefix="/reports", tags=["reports"])
EXPORT_HEADERS = [
    "Period",
    "Audit Area",
    "Checklist Item",
    "Risk",
    "Observation",
    "Auditor Remark",
    "Responsible",
    "Due Date",
    "Status",
]


def report_filter(statement, report_type: str):
    if report_type == "Exception Report":
        return statement.where(Observation.risk.in_(["HIGH", "CRITICAL"]))
    if report_type == "Pending Action Taken Report":
        return statement.where(Observation.status.in_(["PENDING", "REPEATED"]))
    if report_type == "GST Compliance Report":
        return statement.where(
            or_(
                AuditArea.name.ilike("%GST%"),
                AuditArea.name.ilike("%Sales%"),
                AuditArea.name.ilike("%Purchases%"),
            )
        )
    if report_type == "TDS Compliance Report":
        return statement.where(
            or_(AuditArea.name.ilike("%TDS%"), AuditArea.name.ilike("%Expenses%"))
        )
    if report_type == "Management Discussion Report":
        return statement.where(
            or_(
                Observation.risk.in_(["HIGH", "CRITICAL"]),
                Observation.status.in_(["PENDING", "REPEATED"]),
            )
        )
    return statement


def report_rows(db: Session, report: Report) -> list[list[str]]:
    if report.report_type == "Mandatory Completion Report":
        checklist = db.execute(
            select(ChecklistItem, AuditArea.name)
            .join(AuditArea, AuditArea.id == ChecklistItem.audit_area_id)
            .where(
                AuditArea.organization_id == report.organization_id,
                AuditArea.is_active.is_(True),
                ChecklistItem.is_active.is_(True),
                ChecklistItem.is_mandatory.is_(True),
            )
            .order_by(AuditArea.sort_order, ChecklistItem.sort_order)
        ).all()
        observations = {
            item.checklist_item_id: item
            for item in db.scalars(
                select(Observation).where(
                    Observation.entity_id == report.entity_id,
                    Observation.period == report.period,
                    Observation.is_active.is_(True),
                )
            ).all()
        }
        return [
            [
                report.period,
                area_name,
                item.question,
                observations[item.id].risk.value if item.id in observations else "",
                observations[item.id].observation if item.id in observations else "MISSING",
                observations[item.id].remark
                if item.id in observations
                else "Mandatory response incomplete",
                observations[item.id].responsible_person if item.id in observations else "",
                observations[item.id].due_date.isoformat()
                if item.id in observations and observations[item.id].due_date
                else "",
                observations[item.id].status.value if item.id in observations else "Incomplete",
            ]
            for item, area_name in checklist
        ]
    statement = (
        select(Observation, AuditArea.name, ChecklistItem.question)
        .join(AuditArea, AuditArea.id == Observation.audit_area_id)
        .join(ChecklistItem, ChecklistItem.id == Observation.checklist_item_id)
        .where(
            Observation.entity_id == report.entity_id,
            Observation.period == report.period,
            Observation.is_active.is_(True),
        )
    )
    result = db.execute(
        report_filter(statement, report.report_type).order_by(
            AuditArea.sort_order, ChecklistItem.sort_order
        )
    ).all()
    return [
        [
            observation.period,
            area_name,
            question,
            observation.risk.value,
            observation.observation,
            observation.remark,
            observation.responsible_person,
            observation.due_date.isoformat() if observation.due_date else "",
            observation.status.value,
        ]
        for observation, area_name, question in result
    ]


def setting_text(db: Session, organization_id, key: str, fallback: str) -> str:
    setting = db.scalar(
        select(AppSetting).where(
            AppSetting.organization_id == organization_id, AppSetting.key == key
        )
    )
    if not setting:
        return fallback
    value = setting.value
    if isinstance(value, dict):
        value = value.get("value", fallback)
    return str(value or fallback)


def export_name(report: Report, extension: str) -> str:
    base = report.report_number or f"draft-{report.period}-{report.id}"
    safe = "_".join(base.replace("/", "-").split())
    return f"{safe}.{extension}"


def completion(db: Session, organization_id, entity_id, period) -> tuple[int, int]:
    total = (
        db.scalar(
            select(func.count(ChecklistItem.id))
            .join(AuditArea)
            .where(
                AuditArea.organization_id == organization_id,
                AuditArea.is_active.is_(True),
                ChecklistItem.is_active.is_(True),
                ChecklistItem.is_mandatory.is_(True),
            )
        )
        or 0
    )
    completed = (
        db.scalar(
            select(func.count(Observation.id))
            .join(ChecklistItem)
            .where(
                Observation.organization_id == organization_id,
                Observation.entity_id == entity_id,
                Observation.period == period,
                Observation.is_active.is_(True),
                ChecklistItem.is_mandatory.is_(True),
                func.length(func.trim(Observation.observation)) > 0,
            )
        )
        or 0
    )
    return completed, total


@router.get("", response_model=list[ReportOut])
def list_reports(
    user: User = Depends(require_permission("VIEW_REPORT")), db: Session = Depends(get_db)
):
    stmt = select(Report).where(Report.organization_id == user.organization_id)
    allowed = permitted_entity_ids(user)
    if allowed is not None:
        stmt = stmt.where(Report.entity_id.in_(allowed))
    return db.scalars(stmt.order_by(Report.created_at.desc())).all()


@router.post("/generate", response_model=ReportOut, status_code=status.HTTP_201_CREATED)
def generate_report(
    payload: ReportGenerate,
    user: User = Depends(require_permission("GENERATE_REPORT")),
    db: Session = Depends(get_db),
):
    entity = ensure_entity_access(db, payload.entity_id, user)
    statement = (
        select(Observation)
        .join(AuditArea, AuditArea.id == Observation.audit_area_id)
        .where(
            Observation.entity_id == entity.id,
            Observation.period == payload.period,
            Observation.is_active.is_(True),
        )
    )
    observations = db.scalars(report_filter(statement, payload.report_type)).all()
    complete, total = completion(db, user.organization_id, entity.id, payload.period)
    snapshot = {
        "completion": {"completed": complete, "total": total},
        "observation_ids": [str(item.id) for item in observations],
    }
    report = Report(
        organization_id=user.organization_id,
        entity_id=entity.id,
        period=payload.period,
        report_type=payload.report_type,
        status=ReportStatus.GENERATED,
        snapshot=snapshot,
        generated_by_id=user.id,
    )
    db.add(report)
    db.flush()
    record_audit(db, user, "REPORT_GENERATED", "report", report.id, new_value=snapshot)
    db.commit()
    db.refresh(report)
    return report


@router.post("/{report_id}/approve", response_model=ReportOut)
def approve_and_lock(
    report_id: uuid.UUID,
    user: User = Depends(require_permission("APPROVE_REPORT")),
    db: Session = Depends(get_db),
):
    report = db.scalar(select(Report).where(Report.id == report_id).with_for_update())
    if not report:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Report not found")
    ensure_entity_access(db, report.entity_id, user)
    if report.status == ReportStatus.LOCKED or db.scalar(
        select(ReportLock.id).where(
            ReportLock.entity_id == report.entity_id, ReportLock.period == report.period
        )
    ):
        raise HTTPException(status.HTTP_409_CONFLICT, "This entity and period is already locked")
    complete, total = completion(db, user.organization_id, report.entity_id, report.period)
    if total == 0 or complete != total:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            f"Cannot approve: {total - complete} mandatory checklist items are incomplete",
        )
    entity = db.get(Entity, report.entity_id)
    financial_year = setting_text(db, user.organization_id, "financial_year", report.period[:4])
    sequence = db.scalar(
        select(ReportSequence)
        .where(
            ReportSequence.organization_id == user.organization_id,
            ReportSequence.financial_year == financial_year,
        )
        .with_for_update()
    )
    if not sequence:
        sequence = ReportSequence(
            organization_id=user.organization_id, financial_year=financial_year, next_value=1
        )
        db.add(sequence)
        db.flush()
    number = sequence.next_value
    sequence.next_value += 1
    entity_code = "".join(char for char in entity.name.upper() if char.isalnum())[:8]
    prefix = setting_text(db, user.organization_id, "report_prefix", "IA/MWB").strip("/")
    report.report_number = f"{prefix}/{entity_code}/{report.period}/{number:04d}"
    report.status = ReportStatus.LOCKED
    report.approved_by_id = user.id
    report.approved_at = datetime.now(timezone.utc)
    db.add(
        ReportLock(
            report_id=report.id,
            entity_id=report.entity_id,
            period=report.period,
            locked_by_id=user.id,
        )
    )
    db.query(Observation).filter(
        Observation.entity_id == report.entity_id, Observation.period == report.period
    ).update({"locked_at": report.approved_at})
    record_audit(
        db,
        user,
        "REPORT_APPROVED",
        "report",
        report.id,
        new_value={"report_number": report.report_number},
    )
    record_audit(
        db,
        user,
        "REPORT_LOCKED",
        "report",
        report.id,
        new_value={"entity_id": str(report.entity_id), "period": report.period},
    )
    db.commit()
    db.refresh(report)
    return report


@router.get("/{report_id}/export/{export_format}")
def export_report(
    report_id: uuid.UUID,
    export_format: str,
    user: User = Depends(require_permission("VIEW_REPORT")),
    db: Session = Depends(get_db),
):
    report = db.get(Report, report_id)
    if not report:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "Report not found")
    ensure_entity_access(db, report.entity_id, user)
    rows = report_rows(db, report)
    firm_name = setting_text(db, user.organization_id, "firm", "Internal Audit Firm")
    client_name = setting_text(db, user.organization_id, "client", "Client")
    export_format = export_format.lower()
    if export_format == "csv":
        output = StringIO()
        writer = csv.writer(output)
        writer.writerow(EXPORT_HEADERS)
        writer.writerows(rows)
        content = BytesIO(output.getvalue().encode("utf-8-sig"))
        media, extension = "text/csv", "csv"
    elif export_format == "xlsx":
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Audit Report"
        sheet.append([firm_name])
        sheet.append([client_name, report.report_type, report.report_number or "Draft"])
        sheet.append(EXPORT_HEADERS)
        for row in rows:
            sheet.append(row)
        sheet.freeze_panes = "A4"
        sheet.auto_filter.ref = sheet.dimensions
        for cell in sheet[3]:
            cell.font = cell.font.copy(bold=True)
        for column in sheet.columns:
            sheet.column_dimensions[column[0].column_letter].width = min(
                55, max(12, max(len(str(cell.value or "")) for cell in column) + 2)
            )
        content = BytesIO()
        workbook.save(content)
        content.seek(0)
        media, extension = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "xlsx",
        )
    elif export_format == "docx":
        document = WordDocument()
        document.add_heading(report.report_type, 0)
        document.add_paragraph(f"{firm_name}\nInternal Audit Assignment – {client_name}")
        document.add_paragraph(
            f"Period: {report.period}\nReport No: {report.report_number or 'Draft'}"
        )
        table = document.add_table(rows=1, cols=len(EXPORT_HEADERS))
        table.style = "Table Grid"
        for index, heading in enumerate(EXPORT_HEADERS):
            table.rows[0].cells[index].text = heading
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
        content = BytesIO()
        document.save(content)
        content.seek(0)
        media, extension = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    elif export_format == "pdf":
        content = BytesIO()
        styles = getSampleStyleSheet()
        pdf = SimpleDocTemplate(
            content,
            pagesize=landscape(A4),
            leftMargin=10 * mm,
            rightMargin=10 * mm,
            topMargin=10 * mm,
            bottomMargin=10 * mm,
        )
        story = [
            Paragraph(report.report_type, styles["Title"]),
            Paragraph(
                f"{firm_name} – Internal Audit Assignment for {client_name}", styles["Normal"]
            ),
            Paragraph(
                f"Period: {report.period} &nbsp;&nbsp; Report No: {report.report_number or 'Draft'}",
                styles["Normal"],
            ),
            Spacer(1, 5 * mm),
        ]
        pdf_rows = [[Paragraph(value, styles["BodyText"]) for value in EXPORT_HEADERS]] + [
            [Paragraph(str(value), styles["BodyText"]) for value in row] for row in rows
        ]
        table = Table(
            pdf_rows,
            repeatRows=1,
            colWidths=[
                18 * mm,
                25 * mm,
                42 * mm,
                16 * mm,
                52 * mm,
                36 * mm,
                24 * mm,
                20 * mm,
                20 * mm,
            ],
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dfe8ff")),
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("LEFTPADDING", (0, 0), (-1, -1), 3),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        story.append(table)
        pdf.build(story)
        content.seek(0)
        media, extension = "application/pdf", "pdf"
    else:
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY, "Format must be csv, xlsx, docx, or pdf"
        )
    record_audit(
        db, user, "REPORT_EXPORTED", "report", report.id, new_value={"format": export_format}
    )
    db.commit()
    return StreamingResponse(
        content,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{export_name(report, extension)}"'},
    )
